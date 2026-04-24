import io
from docx import Document

def create_worksheet_docx(title, questions, answers=None, include_answers=False):
    doc = Document()
    doc.add_heading(title, 0)
    
    for i, q in enumerate(questions):
        clean_q = str(q).strip()
        # 如果题目不以数字开头，自动加上编号
        if not clean_q[0].isdigit():
            clean_q = f"{i+1}. {clean_q}"
        doc.add_paragraph(clean_q)
    
    if include_answers and answers:
        doc.add_heading("答案", level=1)
        for i, a in enumerate(answers):
            doc.add_paragraph(f"{i+1}. {str(a).strip()}")
    
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()